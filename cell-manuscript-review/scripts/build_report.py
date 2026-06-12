#!/usr/bin/env python3
"""Build the Cell submission-readiness review .docx from a JSON findings file.

A python-docx reimplementation of build_report.js for environments without
Node.js (e.g. Windows). It accepts the same JSON schema (see
references/findings-schema.example.json and references/report-format.md) and
matches the layout: title block, executive summary, at-a-glance table,
severity-grouped detail entries, separate-upload list, and checklist.

Usage:
  python build_report.py <findings.json> <output.docx>
"""
import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if len(sys.argv) < 3:
    sys.exit("Usage: python build_report.py <findings.json> <output.docx>")

data = json.loads(open(sys.argv[1], encoding="utf-8").read())
out_path = sys.argv[2]

SEV_COLORS = {"CRITICAL": "C00000", "MAJOR": "E97132", "MINOR": "BF9000", "INFO": "548235"}
ACCENT = "1F4E79"
ACTION_BLUE = "1F4E79"
MARK = {"critical": ("Ⓒ", "C00000"), "major": ("Ⓜ", "E97132"), "separate": ("Ⓢ", "548235")}

doc = Document()

# base style + Letter page with narrow margins
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
sec = doc.sections[0]
sec.page_width, sec.page_height = Twips(12240), Twips(15840)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Twips(1080)

for hname, sz, col in (("Heading 1", 15, ACCENT), ("Heading 2", 13, "2E75B6")):
    st = doc.styles[hname]
    st.font.name = "Arial"
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(col)


def run(p, text, *, bold=False, italic=False, color=None, size=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    return r


def para(*, before=6, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def bullet(text_runs, *, marker=None, marker_color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if marker:
        run(p, marker + "  ", bold=True, color=marker_color)
    if isinstance(text_runs, str):
        run(p, text_runs)
    else:
        for t, kw in text_runs:
            run(p, t, **kw)
    return p


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=None, align=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run(p, text, bold=bold, color=color, size=size)


# ---------- title block ----------
rev = data.get("revision_number", 1)
title_text = f"Cell Submission Readiness Review — Revision {rev}" if rev and rev > 1 else "Cell Submission Readiness Review"
p = para(after=5)
run(p, title_text, bold=True, color=ACCENT, size=20)

p = para(after=12)
run(p, "Manuscript: ", bold=True, size=11)
run(p, data.get("manuscript_title", "(untitled)"), italic=True, size=11)

p = para(after=12)
run(p, "Reference standards: ", bold=True, size=10)
run(p, "Cell Final File Requirements (CELLFFC.pdf); STAR Methods Article Template; Cell Information for Authors.", size=10)

# ---------- executive summary ----------
es = data.get("executive_summary")
if es:
    doc.add_heading("Executive summary", level=1)
    for key in ("good_state", "blockers"):
        if es.get(key):
            run(para(), es[key])
    for b in es.get("blocker_list", []):
        bullet(b)
    if es.get("revision_note"):
        run(para(), es["revision_note"])

# ---------- fixed since previous ----------
fixed = data.get("fixed_since_previous") or []
if fixed:
    doc.add_heading("What was fixed since the previous review", level=1)
    run(para(), "These items from the previous review are now resolved:")
    for item in fixed:
        bullet(item)

# ---------- at a glance ----------
findings = data.get("findings", [])
sev_order = {"CRITICAL": 0, "MAJOR": 1, "MINOR": 2, "INFO": 3}
if findings:
    doc.add_heading("Issues at a glance", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    widths = (Twips(1200), Twips(5160), Twips(2880))
    hdr = table.rows[0].cells
    for c, label, al in ((hdr[0], "Severity", WD_ALIGN_PARAGRAPH.CENTER),
                          (hdr[1], "Issue", None), (hdr[2], "Location", None)):
        shade_cell(c, ACCENT)
        set_cell_text(c, label, bold=True, color="FFFFFF", align=al, size=10)
    for f in sorted(findings, key=lambda x: sev_order.get(x["severity"], 9)):
        cells = table.add_row().cells
        shade_cell(cells[0], SEV_COLORS.get(f["severity"], "808080"))
        set_cell_text(cells[0], f["severity"], bold=True, color="FFFFFF",
                      align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        set_cell_text(cells[1], f.get("summary") or f["title"], size=10)
        set_cell_text(cells[2], f.get("location", ""), size=10)
    for row in table.rows:
        for c, w in zip(row.cells, widths):
            c.width = w

    # ---------- detailed entries grouped by severity ----------
    def detail(f):
        color = SEV_COLORS.get(f["severity"], "404040")
        p = para(before=12, after=4)
        run(p, f["severity"] + "  ", bold=True, color=color, size=10)
        run(p, f["title"], bold=True, size=12)
        p = para(before=0, after=4)
        run(p, "Location: ", bold=True, italic=True, size=10)
        run(p, f.get("location", ""), italic=True, size=10)
        p = para(before=0, after=3)
        run(p, "Issue: ", bold=True, size=10)
        run(p, f.get("issue", ""), size=10)
        p = para(before=0, after=4)
        run(p, "Action: ", bold=True, color=ACTION_BLUE, size=10)
        run(p, f.get("action", ""), size=10)

    groups = [("CRITICAL", "Critical issues",
               "These items should be addressed before the file is uploaded."),
              ("MAJOR", "Major formatting and structural issues", None),
              ("MINOR", "Minor issues worth addressing", None)]
    for sev, heading, lead in groups:
        items = [f for f in findings if f["severity"] == sev]
        if not items:
            continue
        doc.add_heading(heading, level=1)
        if lead:
            run(para(), lead)
        for f in items:
            detail(f)

# ---------- separate uploads ----------
ups = data.get("separate_uploads") or []
if ups:
    doc.add_heading("Items to prepare and upload separately", level=1)
    run(para(), "These items belong outside the main manuscript file.")
    for item in ups:
        doc.add_heading(item["title"], level=2)
        if item.get("description"):
            run(para(), item["description"])

# ---------- checklist ----------
cl = data.get("checklist") or []
if cl:
    doc.add_heading("Submission readiness checklist", level=1)
    run(para(), "Priority-ordered to-do list. Ⓒ = critical/submission-blocking; "
                "Ⓜ = major formatting fix; Ⓢ = separate upload.")
    for item in cl:
        m, c = MARK.get(item.get("priority"), ("•", "404040"))
        bullet(item["text"], marker=m, marker_color=c)

# ---------- end marker ----------
p = para(before=24, after=6)
run(p, "End of review.", italic=True, color="808080")

doc.save(out_path)
print(f"Wrote {out_path}: {len(findings)} findings")
